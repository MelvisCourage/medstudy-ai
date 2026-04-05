from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def export_to_docx(qa_pairs: list[tuple[str, str]], subject: str) -> bytes:
    """Export Q&A pairs to a formatted Word document. Returns bytes."""
    doc = Document()

    # Title
    title = doc.add_heading(f"{subject} — Exam Q&A", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    doc.add_paragraph(f"Total Questions: {len(qa_pairs)}")
    doc.add_paragraph("")

    for i, (question, answer) in enumerate(qa_pairs, 1):
        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{i}. {question}")
        q_run.bold = True
        q_run.font.size = Pt(11)
        q_run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

        # Answer
        a_para = doc.add_paragraph()
        a_run = a_para.add_run(f"A: {answer}")
        a_run.font.size = Pt(10.5)
        a_para.paragraph_format.left_indent = Inches(0.3)

        doc.add_paragraph("")  # spacing

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
